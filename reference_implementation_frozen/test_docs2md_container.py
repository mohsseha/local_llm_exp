#!/usr/bin/env python3
"""
test_docs2md_container.py

Purpose:
Generates a deterministic and challenging set of test files in a deep
directory structure to validate the robustness of the docs2md container.
This script is designed to be simple, self-contained, and run once.

It creates a directory structure inside './data' and populates it with
standard business documents and a "gauntlet" of problematic .eml files
that target common parsing and conversion edge cases.

Required Dependencies (must be installed by the user):
- pandas
- openpyxl
- python-docx
- reportlab
"""

import email
from email.message import EmailMessage
from email.header import Header
from pathlib import Path
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# --- Constants and Hardcoded Paths ---

# The single, hardcoded path where all test data will be generated.
BASE_PATH = Path("./data/FY2023/Q4/Investor_Relations/")

# A standard block of text to be used in generated documents.
SAMPLE_TEXT_CONTENT = (
    "Project Nightingale: Q4 Update\n\n"
    "This document provides a summary of the key performance indicators (KPIs) "
    "for Project Nightingale in the fourth quarter of Fiscal Year 2023. "
    "The primary objective was to stabilize the platform and reduce technical debt.\n\n"
    "Key Achievements:\n"
    "- Platform uptime increased to 99.98%.\n"
    "- Critical bug count reduced by 45%.\n"
    "- Deployed new data analytics module.\n\n"
    "Challenges:\n"
    "Integration with the legacy CRM system remains a bottleneck. "
    "Further resource allocation is required for Q1 2024.\n"
    "The Euro (€) is our primary currency for reporting."
)


def create_directory_structure():
    """Creates the deep, nested directory structure for the test files."""
    print(f"Ensuring directory exists: {BASE_PATH}")
    BASE_PATH.mkdir(parents=True, exist_ok=True)


def generate_standard_documents():
    """
    Generates the '20%' of standard business documents (.txt, .docx, .pdf, .xlsx).
    These files provide a baseline for testing non-EML conversion.
    """
    print("Generating standard business documents...")

    # 1. Create a simple .txt file
    txt_path = BASE_PATH / "report.txt"
    txt_path.write_text(SAMPLE_TEXT_CONTENT, encoding='utf-8')
    print(f"- Created: {txt_path}")

    # 2. Create a simple .docx file
    docx_path = BASE_PATH / "report.docx"
    doc = Document()
    doc.add_heading("Project Nightingale: Q4 Update", level=1)
    doc.add_paragraph(SAMPLE_TEXT_CONTENT)
    doc.save(docx_path)
    print(f"- Created: {docx_path}")

    # 3. Create a simple .pdf file
    pdf_path = BASE_PATH / "report.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    textobject = c.beginText(40, 750)
    for line in SAMPLE_TEXT_CONTENT.split('\n'):
        textobject.textLine(line)
    c.drawText(textobject)
    c.save()
    print(f"- Created: {pdf_path}")

    # 4. Create a multi-sheet .xlsx file
    xlsx_path = BASE_PATH / "financials.xlsx"
    with pd.ExcelWriter(xlsx_path, engine='openpyxl') as writer:
        df1_data = {
            'Region': ['North America', 'Europe', 'Asia'],
            'Revenue (M \u20ac)': [10.5, 8.2, 5.7],
            'Growth (%)': [5.2, 3.1, 10.8]
        }
        df1 = pd.DataFrame(df1_data)
        df1.to_excel(writer, sheet_name='Regional Summary', index=False)

        df2_data = {
            'Expense Category': ['R&D', 'Marketing', 'G&A'],
            'Budget (M \u20ac)': [4.1, 2.5, 1.8],
            'Actual (M \u20ac)': [4.5, 2.4, 1.9]
        }
        df2 = pd.DataFrame(df2_data)
        df2.to_excel(writer, sheet_name='Expense Report', index=False)
    print(f"- Created: {xlsx_path}")


def generate_eml_gauntlet():
    """
    Generates the '80%' of difficult .eml files to test edge cases.
    Each file is crafted to trigger a specific potential failure mode.
    """
    print("Generating .eml gauntlet...")

    # --- Helper to save .eml files ---
    def save_eml(msg: EmailMessage, filename: str):
        path = BASE_PATH / filename
        with open(path, 'wb') as f:
            f.write(msg.as_bytes())
        print(f"- Created: {path}")

    # 1. Malformed Date Header
    msg = EmailMessage()
    msg['From'] = "sender@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "Invalid Date Test"
    msg['Date'] = "Sometime last week"  # Invalid date format
    msg.set_content("This email has a date that is not RFC 2822 compliant.")
    save_eml(msg, "malformed_date.eml")

    # 2. Encoding Mismatch (Header says UTF-8, body is Windows-1252)
    msg = EmailMessage()
    msg['From'] = "sender@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "Encoding Mismatch"
    # The content contains the Euro symbol (€), which has a different byte representation
    # in UTF-8 vs. Windows-1252. We encode it as Windows-1252 bytes.
    euro_content = "This invoice is for 100€."
    # The headers incorrectly claim the content is UTF-8. A robust parser must handle this.
    msg.set_content(euro_content, subtype='plain', charset='utf-8')
    # Now, we manually replace the payload with the incorrectly encoded bytes.
    del msg['Content-Transfer-Encoding']
    msg.set_payload(euro_content.encode('windows-1252'))
    save_eml(msg, "encoding_mismatch.eml")

    # 3. HTML-Only Content (No text/plain part)
    msg = EmailMessage()
    msg['From'] = "newsletter@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "HTML Only Newsletter"
    html_content = "<html><body><h1>Q4 Results</h1><p>We saw a <b>significant</b> increase in growth.</p></body></html>"
    msg.set_content(html_content, subtype='html')
    save_eml(msg, "html_only.eml")

    # 4. Mock winmail.dat (TNEF) attachment from old Outlook clients
    msg = EmailMessage()
    msg['From'] = "corporate.outlook@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "Meeting Follow-up with TNEF"
    msg.set_content("Please see the attached notes.")
    # A small, fake TNEF payload. The key is the content type.
    tnef_payload = b'\x78\x9c\x63\x64\x62\x66\x61\x65\x03\x00\x00\x0b\x00\x04'
    msg.add_attachment(tnef_payload, maintype='application', subtype='ms-tnef', filename='winmail.dat')
    save_eml(msg, "winmail_attachment.eml")

    # 5 & 6. A two-part email thread
    # Part 1: The original email
    msg1 = EmailMessage()
    msg1['From'] = "analyst@example.com"
    msg1['To'] = "cfo@example.com"
    msg1['Subject'] = "Series D Funding Inquiry"
    msg1['Date'] = "Mon, 1 Dec 2023 10:00:00 -0500"
    msg1['Message-ID'] = "<part1.12345@example.com>"
    msg1.set_content("Could you provide the latest cap table for our Series D model?")
    save_eml(msg1, "thread_part1_of_2.eml")

    # Part 2: The reply
    msg2 = EmailMessage()
    msg2['From'] = "cfo@example.com"
    msg2['To'] = "analyst@example.com"
    msg2['Subject'] = "Re: Series D Funding Inquiry"
    msg2['Date'] = "Mon, 1 Dec 2023 11:30:00 -0500"
    msg2['Message-ID'] = "<part2.67890@example.com>"
    msg2['In-Reply-To'] = "<part1.12345@example.com>"
    msg2['References'] = "<part1.12345@example.com>"
    msg2.set_content("Attached. Let me know if you have questions.")
    save_eml(msg2, "thread_part2_of_2.eml")

    # 7. Calendar Invite (.ics attachment)
    msg = EmailMessage()
    msg['From'] = "scheduler@example.com"
    msg['To'] = "board@example.com"
    msg['Subject'] = "Board Meeting Q4"
    
    # To attach parts, the main message must be multipart.
    msg.make_mixed()
    
    # Attach the plain text body part
    msg.attach(EmailMessage())
    msg.get_payload()[0].set_content("Please see the attached meeting invitation.")

    ics_content = (
        "BEGIN:VCALENDAR\n"
        "VERSION:2.0\n"
        "BEGIN:VEVENT\n"
        "SUMMARY:Q4 Board Meeting\n"
        "DTSTART:20231215T130000Z\n"
        "DTEND:20231215T150000Z\n"
        "END:VEVENT\n"
        "END:VCALENDAR"
    )
    # Now, attach the calendar part
    ics_part = EmailMessage()
    ics_part.set_content(ics_content)
    ics_part.replace_header('Content-Type', 'text/calendar; method=REQUEST; name="invite.ics"')
    ics_part.add_header('Content-Disposition', 'attachment', filename='invite.ics')
    msg.attach(ics_part)
    save_eml(msg, "calendar_invite.eml")

    # 8. Corrupted Body (contains null bytes)
    msg = EmailMessage()
    msg['From'] = "corrupted.system@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "System Alert - Corrupted"
    # A normal string with null bytes inserted.
    corrupted_payload = b'This is a system alert.\x00\x00\x00The data feed is down.'
    msg.set_payload(corrupted_payload)
    del msg['Content-Transfer-Encoding'] # Let the library figure it out
    save_eml(msg, "corrupted_body.eml")

    # --- NEW CORNER CASES ---

    # 9. No Message-ID Header
    msg = EmailMessage()
    msg['From'] = "no-id-sender@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "Email without a Message-ID"
    msg['Date'] = "Tue, 2 Dec 2023 11:00:00 -0500"
    msg.set_content("This email is used to test threading fallback logic.")
    save_eml(msg, "no_message_id.eml")

    # 10. Extremely Long Subject
    msg = EmailMessage()
    msg['From'] = "sender@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "SPAM SPAM SPAM " + ("A" * 2000)
    msg.set_content("This email has an excessively long subject line.")
    save_eml(msg, "extremely_long_subject.eml")

    # 11. Forwarded Email (message/rfc822 attachment)
    msg = EmailMessage()
    msg['From'] = "cfo@example.com"
    msg['To'] = "legal@example.com"
    msg['Subject'] = "Fwd: Series D Funding Inquiry"
    msg.set_content("FYI - see the original inquiry below.")
    # We'll use the bytes of the first thread email as the attachment
    original_email_bytes = msg1.as_bytes()
    msg.add_attachment(original_email_bytes, maintype='message', subtype='rfc822', filename="original_inquiry.eml")
    save_eml(msg, "forwarded_email.eml")

    # 12. Quoted-Printable Body
    msg = EmailMessage()
    msg['From'] = "international.sender@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "Notes with special characters: éàçü"
    # This content is likely to be encoded as quoted-printable
    qp_content = (
        "This is a line that is designed to be longer than the typical 76-character limit "
        "to see how the line wrapping is handled.\n"
        "It also includes special characters like résumé and façade.\n"
        "Final line."
    )
    # The library will automatically choose quoted-printable if we set the CTE
    msg.set_content(qp_content, cte='quoted-printable')
    save_eml(msg, "quoted_printable_body.eml")

    # 13. Attachment with no filename
    msg = EmailMessage()
    msg['From'] = "sender@example.com"
    msg['To'] = "recipient@example.com"
    msg['Subject'] = "Attachment with no filename"
    msg.make_mixed() # Convert to multipart before attaching

    # Attach the plain text body part
    msg.attach(EmailMessage())
    msg.get_payload()[0].set_content("The attached log file is missing a filename in the header.")

    # The key is the Content-Disposition header lacks the 'filename' parameter
    log_part = EmailMessage()
    log_part.set_content("INFO: System startup complete.")
    log_part.replace_header('Content-Type', 'text/plain')
    log_part.add_header('Content-Disposition', 'attachment') # No filename!
    msg.attach(log_part)
    save_eml(msg, "attachment_no_filename.eml")

    # --- NIGHTMARE CORNER CASES ---

    # 14. Mixed Language, Emojis, and Encodings
    msg = EmailMessage()
    msg['From'] = "global-team@example.com"
    msg['To'] = "cfo@example.com"
    # Subject contains multiple languages and a multi-codepoint emoji
    subject_text = "🚀 Project Phoenix Update 🧑‍🚀 (م更新)"
    msg['Subject'] = subject_text # The library handles the encoding automatically
    
    text_plain = "Hello Team,\nPlease find the project update attached."
    
    html_body = """
    <html>
      <head>
        <meta charset="UTF-8">
      </head>
      <body>
        <p>Hello Team,</p>
        <p>Please find the project update attached.</p>
        <p dir="rtl">مرحبا بالعالم</p>
        <p>你好世界</p>
        <p>And here is a complex emoji: 🧑‍🚀</p>
      </body>
    </html>
    """
    msg.set_content(text_plain, subtype='plain', charset='utf-8')
    msg.add_alternative(html_body, subtype='html', charset='utf-8')
    save_eml(msg, "mixed_language_emojis_and_encoding.eml")

    # 15. Attachment with Exotic Filename (RFC 2231 encoding)
    msg = EmailMessage()
    msg['From'] = "international-hr@example.com"
    msg['To'] = "cfo@example.com"
    msg['Subject'] = "Employee CV"
    msg.set_content("Please review the attached CV.")
    # This filename requires complex encoding that many parsers fail to handle.
    exotic_filename = "résumé-français-日本語-report.docx"
    # A tiny mock docx payload
    docx_payload = b'PK\x05\x06\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00'
    msg.add_attachment(docx_payload, maintype='application', 
                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                       filename=exotic_filename)
    save_eml(msg, "attachment_with_exotic_filename.eml")

    # 16. Deeply Nested and Related Multipart
    msg = EmailMessage()
    msg['From'] = "marketing-automation@example.com"
    msg['To'] = "cfo@example.com"
    msg['Subject'] = "Weekly Marketing Newsletter"
    
    msg.make_mixed() # Level 1: container for body and regular attachments
    
    # Create the body part, which is itself multipart
    body_part = EmailMessage()
    body_part.make_alternative() # Level 2: container for plain and html versions
    msg.attach(body_part)

    # Plain text version
    body_part.attach(EmailMessage())
    body_part.get_payload()[0].set_content("This is the newsletter. Please enable HTML to see images.")

    # HTML version, which is *also* multipart to hold the inline image
    html_part_container = EmailMessage()
    html_part_container.make_related() # Level 3: container for HTML and its inline images
    body_part.attach(html_part_container)

    # The actual HTML
    html_part_container.attach(EmailMessage())
    html_part_container.get_payload()[0].set_content(
        '<html><body><p>Check out our new logo!</p><img src="cid:logo_image"></body></html>',
        subtype='html'
    )

    # The inline image data
    # A tiny 1x1 transparent PNG payload
    png_payload = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82'
    inline_image_part = EmailMessage()
    inline_image_part.set_content(png_payload, maintype='image', subtype='png')
    inline_image_part.add_header('Content-ID', '<logo_image>')
    inline_image_part.add_header('Content-Disposition', 'inline', filename='logo.png')
    html_part_container.attach(inline_image_part)

    # Finally, add a regular PDF attachment at the top level
    msg.add_attachment(b'%PDF-1.0\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n',
                       maintype='application', subtype='pdf', filename='newsletter_promo.pdf')
    save_eml(msg, "deeply_nested_and_related_multipart.eml")

    # 17. Header Injection Attack Attempt
    # The python email library prevents this, so we construct the raw text manually.
    raw_eml_content = (
        'From: "Valid Name <valid@example.com>\r\nBcc: hidden@example.com"\r\n'
        'To: recipient@example.com\r\n'
        'Subject: Header Injection Test\r\n'
        '\r\n'
        'This email tests for header injection vulnerabilities.'
    )
    eml_path = BASE_PATH / "header_injection_attack.eml"
    with open(eml_path, 'w', encoding='ascii') as f:
        f.write(raw_eml_content)
    print(f"- Created: {eml_path}")


def main():
    """Main execution function."""
    print("Starting test data generation...")
    create_directory_structure()
    generate_standard_documents()
    generate_eml_gauntlet()
    print("\nTest data generation complete.")
    print(f"All files have been created in: {BASE_PATH.resolve()}")


if __name__ == "__main__":
    main()
